from argparse import Action
from urllib import response
from rest_framework.viewsets import GenericViewSet
from rest_framework import mixins, viewsets
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework import serializers

from content.models import Book,Kind,Form,Matter,Author,PublishingHouse
from content.serializers import BookSerializers,KindSerializer,FormSerializer,MatterSerializer,AuthorSerializer,PublishingHouseSerializer

from django.contrib.auth import authenticate,login, logout
from rest_framework.permissions import IsAuthenticated
from rest_framework.permissions import BasePermission

from django.db.models import Avg,Count,Min,Max

import openpyxl
from openpyxl.styles import Font
from django.http import HttpResponse
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework import status
from io import BytesIO
from openpyxl import Workbook
from docx import Document
import pyotp
from django.core.cache import cache 






from .models import UserProfile, User
from .serializers import UserProfileSerializer

class BaseUserViewSet(
        mixins.ListModelMixin,
        mixins.CreateModelMixin,
        mixins.UpdateModelMixin,
        mixins.RetrieveModelMixin,
        mixins.DestroyModelMixin,
        GenericViewSet
    ):

    def get_queryset(self):
        qs = super().get_queryset()
        user = self.request.user
        if user.is_superuser:
            # Если есть фильтрация по пользователю, то применяем её
            user_id = self.request.query_params.get("user_id")
            if user_id:
                qs = qs.filter(user_id=user_id)
        else:
            # Обычные пользователи видят только свои данные
            qs = qs.filter(user=user)
        return qs



class KindsViewset(BaseUserViewSet):
    queryset=Kind.objects.all()
    serializer_class=KindSerializer
    permission_classes = [IsAuthenticated]


    class StatsSerializer(serializers.Serializer):
        count = serializers.IntegerField()
        avg = serializers.FloatField()
        max = serializers.IntegerField()
        min = serializers.IntegerField()

    @action(detail=False, methods=["GET"], url_path="stats")
    def get_stats(self, request, *args, **kwargs):
        stats = Kind.objects.aggregate(
            count=Count("*"),
            avg=Avg("id"),
            max=Max("id"),
            min=Min("id"),
        )
        serializer = self.StatsSerializer(instance=stats)

        return Response(serializer.data)


class FormsViewset(BaseUserViewSet):
    queryset=Form.objects.all()
    serializer_class=FormSerializer

    class StatsSerializer(serializers.Serializer):
        count=serializers.IntegerField()
        avg=serializers.FloatField()
        max=serializers.IntegerField()
        min=serializers.IntegerField()

    @action(detail=False,methods=["GET"],url_path="stats")
    def get_stats(self,request,*args,**kwargs):
        stats=Form.objects.aggregate(
            count=Count("*"),
            avg=Avg("id"),
            max=Max("id"),
            min=Min("id"),
        )
        serializer=self.StatsSerializer(instance=stats)

        return Response(serializer.data)

class MattersViewset(BaseUserViewSet):
    queryset=Matter.objects.all()
    serializer_class=MatterSerializer

    class StatsSerializer(serializers.Serializer):
        count=serializers.IntegerField()
        avg=serializers.FloatField()
        max=serializers.IntegerField()
        min=serializers.IntegerField()

    @action(detail=False,methods=["GET"],url_path="stats")
    def get_stats(self,request,*args,**kwargs):
        stats=Matter.objects.aggregate(
            count=Count("*"),
            avg=Avg("id"),
            max=Max("id"),
            min=Min("id"),
        )
        serializer=self.StatsSerializer(instance=stats)

        return Response(serializer.data)

class AuthorsViewset(BaseUserViewSet):
    queryset=Author.objects.all()
    serializer_class=AuthorSerializer

    class StatsSerializer(serializers.Serializer):
        count=serializers.IntegerField()
        avg=serializers.FloatField()
        max=serializers.IntegerField()
        min=serializers.IntegerField()

    @action(detail=False,methods=["GET"],url_path="stats")
    def get_stats(self,request,*args,**kwargs):
        stats=Author.objects.aggregate(
            count=Count("*"),
            avg=Avg("id"),
            max=Max("id"),
            min=Min("id"),
        )
        serializer=self.StatsSerializer(instance=stats)

        return Response(serializer.data)

class PublishingHousesViewset(BaseUserViewSet):
    queryset=PublishingHouse.objects.all()
    serializer_class=PublishingHouseSerializer
 

    class StatsSerializer(serializers.Serializer):
        count=serializers.IntegerField()
        avg=serializers.FloatField()
        max=serializers.IntegerField()
        min=serializers.IntegerField()

    @action(detail=False,methods=["GET"],url_path="stats")
    def get_stats(self,request,*args,**kwargs):
        stats=PublishingHouse.objects.aggregate(
            count=Count("*"),
            avg=Avg("year_of_foundation"),
            max=Max("year_of_foundation"),
            min=Min("year_of_foundation"),

        )
        serializer=self.StatsSerializer(instance=stats)

        return Response(serializer.data)


class BooksViewset(BaseUserViewSet):
    queryset=Book.objects.all()
    serializer_class=BookSerializers

    class StatsSerializer(serializers.Serializer):
        count=serializers.IntegerField()
        avg=serializers.FloatField()
        max=serializers.IntegerField()
        min=serializers.IntegerField()

    @action(detail=False,methods=["GET"],url_path="stats")
    def get_stats(self,request,*args,**kwargs):
        stats=Book.objects.aggregate(
            count=Count("*"),
            avg=Avg("year_of_publication"),
            max=Max("year_of_publication"),
            min=Min("year_of_publication"),
        )
        serializer=self.StatsSerializer(instance=stats)

        return Response(serializer.data)
    

    @action(detail=False, methods=["GET"], url_path="export-excel")
    def export_to_excel(self, request, *args, **kwargs):
        # Создаем Excel-файл
        workbook = Workbook()  # Теперь импорт этого класса есть
        sheet = workbook.active
        sheet.title = "Books"

        # Добавляем заголовки
        headers = ["ID", "Title", "Year", "Author", "Matter", "Form","Kind","Publishing house"]
        for col_num, header in enumerate(headers, start=1):
            sheet.cell(row=1, column=col_num, value=header)

        # Добавляем данные
        books = self.queryset.select_related("author_FK", "matter_FK","form_FK","kind_FK","publishingHouse_FK")
        for row_num, book in enumerate(books, start=2):
            sheet.cell(row=row_num, column=1, value=book.id)
            sheet.cell(row=row_num, column=2, value=book.title)
            sheet.cell(row=row_num, column=3, value=book.year_of_publication)
            sheet.cell(row=row_num, column=4, value= (book.author_FK.name[0]+"."+book.author_FK.surname[0]+". "+ book.author_FK.patronymic) if book.author_FK else "")
            sheet.cell(row=row_num, column=5, value=book.matter_FK.title if book.matter_FK else "")
            sheet.cell(row=row_num, column=6, value=book.form_FK.name if book.form_FK else "")
            sheet.cell(row=row_num, column=7, value=book.kind_FK.name if book.kind_FK else "")
            sheet.cell(row=row_num, column=8, value=book.publishingHouse_FK.name if book.publishingHouse_FK else "")

        # Сохранение и отправка файла
        buffer = BytesIO()
        workbook.save(buffer)
        buffer.seek(0)
        response = HttpResponse(
            buffer,
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
        response["Content-Disposition"] = 'attachment; filename="books.xlsx"'
        return response

    @action(detail=False, methods=["GET"], url_path="export-word")
    def export_to_word(self, request, *args, **kwargs):
        # Создаем Word-документ
        document = Document()
        document.add_heading("Book List", level=1)

        # Добавляем данные в таблицу
        books = self.queryset.select_related("author_FK", "matter_FK","form_FK","kind_FK","publishingHouse_FK")
        table = document.add_table(rows=1, cols=8)
        headers = ["ID", "Title", "Year", "Author", "Matter", "Form","Kind","Publishing house"]
        for i, header in enumerate(headers):
            table.cell(0, i).text = header

        for book in books:
            row = table.add_row().cells
            row[0].text = str(book.id)
            row[1].text = book.title
            row[2].text = str(book.year_of_publication)
            row[3].text = (book.author_FK.name[0]+"."+book.author_FK.surname[0]+". "+ book.author_FK.patronymic) if book.author_FK else ""
            row[4].text = book.matter_FK.title if book.matter_FK else ""
            row[5].text = book.form_FK.name if book.form_FK else ""
            row[6].text = book.kind_FK.name if book.kind_FK else ""
            row[7].text = book.publishingHouse_FK.name if book.publishingHouse_FK else ""

        # Сохранение и отправка файла
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        response = HttpResponse(
            buffer,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = 'attachment; filename="books.docx"'
        return response



class BaseUserViewSet(
        mixins.ListModelMixin,
        mixins.CreateModelMixin,
        mixins.UpdateModelMixin,
        mixins.RetrieveModelMixin,
        mixins.DestroyModelMixin,
        GenericViewSet
    ):

    def get_queryset(self):
        qs = super().get_queryset()
        user = self.request.user
        if user.is_superuser:
            # Если есть фильтрация по пользователю, то применяем её
            user_id = self.request.query_params.get("user_id")
            if user_id:
                qs = qs.filter(user_id=user_id)
        else:
            # Обычные пользователи видят только свои данные
            qs = qs.filter(user=user)
        return qs
    

class UserViewSet(GenericViewSet):
    @action(url_path="info", methods=["GET"], detail=False)
    def get_info(self, request, *args, **kwargs):
        data = {
            "is_authenticated": request.user.is_authenticated
        }
        if request.user.is_authenticated:
            data.update({
                "username": request.user.username,
                "user_id": request.user.id
            })
        return Response(data)

    @action(url_path="login", methods=["POST"], detail=False)
    def login(self, request, *args, **kwargs):
        username = request.data.get("user")
        password = request.data.get("password")

        user = authenticate(request, username=username, password=password)
        if user:
            login(request, user)
        return Response({})

    @action(url_path="logout", methods=["POST"], detail=False)
    def logout(self, request, *args, **kwargs):
        logout(request)
        return Response({})



class UserProfileViewSet(viewsets.GenericViewSet):
    permission_classes = [IsAuthenticated]

    class OTPSerializer(serializers.Serializer):
        key = serializers.CharField()

    class OTPRequired(BasePermission):
        def has_permission(self, request, view):
            return bool(request.user and cache.get('otp_good', False))

    @action(detail=False, url_path="check-login", methods=['GET'], permission_classes=[])
    def get_check_login(self, request, *args, **kwargs):
        return Response({
            'is_authenticated': self.request.user.is_authenticated
        })
    
    @action(detail=False, url_path="login", methods=['GET'], permission_classes=[])
    def use_login(self, request, *args, **kwargs):
        user = authenticate(username='username', password='pass')
        if user:
            login(request, user)
        return Response({
            'is_authenticated': bool(user)
        })

    @action(detail=False, url_path='otp-login', methods=['POST'], serializer_class=OTPSerializer)
    def otp_login(self, *args, **kwargs):
        totp = pyotp.TOTP(self.request.user.userprofile.opt_key)
        
        serializer = self.get_serializer(data=self.request.data)
        serializer.is_valid(raise_exception=True)

        success = False
        if totp.now() == serializer.validated_data['key']:
            cache.set('otp_good', True, 600)  # Устанавливаем время жизни флага на 10 минут
            success = True

        return Response({
            'success': success
        })
    
    @action(detail=False, url_path='otp-status')
    def get_otp_status(self, *args, **kwargs):
        otp_good = cache.get('otp_good', False)
        return Response({
            'otp_good': otp_good
        })
    
    @action(detail=False, url_path='otp-required', permission_classes=[OTPRequired])
    def page_with_otp_required(self, *args, **kwargs):
        return Response({
            'success': True
        })

    @action(detail=False, url_path="edit-object", methods=['POST'], permission_classes=[OTPRequired])
    def edit_object(self, request, *args, **kwargs):
        # Логика редактирования объекта
        return Response({
            'success': True
        })
    
class UserProfileViewSet(viewsets.GenericViewSet):
    permission_classes = [IsAuthenticated]

    class OTPSerializer(serializers.Serializer):
        key = serializers.CharField()

    class OTPRequired(BasePermission):
        def has_permission(self, request, view):
            return bool(request.user and cache.get('otp_good', False))

    @action(detail=False, url_path="check-login", methods=['GET'], permission_classes=[])
    def get_check_login(self, request, *args, **kwargs):
        return Response({
            'is_authenticated': self.request.user.is_authenticated
        })

    @action(detail=False, url_path="login", methods=['POST'], permission_classes=[])
    def use_login(self, request, *args, **kwargs):
        user = authenticate(username=request.data.get('user'), password=request.data.get('password'))
        if user:
            login(request, user)
            try:
               UserProfile.objects.get_or_create(user=user)
            except:
                 return Response({
                    'is_authenticated': False
                    })
        return Response({
            'is_authenticated': bool(user)
        })
    @action(detail=False, url_path='get-opt-key', methods=['GET'])
    def get_opt_key(self,request):
        user=request.user
        if user.is_authenticated:
            try:
                user_profile=UserProfile.objects.get(user=user)
                serializer=UserProfileSerializer(user_profile)
                return Response(serializer.data)
            except UserProfile.DoesNotExist:
                return Response({
                    "error": "UserProfile not found"
                }, status=400)
        else:
             return Response({
                "error": "unauthorised"
            }, status=401)
    @action(detail=False, url_path='otp-login', methods=['POST'], serializer_class=OTPSerializer)
    def otp_login(self, *args, **kwargs):
        totp = pyotp.TOTP(self.request.user.userprofile.opt_key)

        serializer = self.get_serializer(data=self.request.data)
        serializer.is_valid(raise_exception=True)

        success = False
        if totp.now() == serializer.validated_data['key']:
            cache.set('otp_good', True, 600)  # Устанавливаем время жизни флага на 10 минут
            success = True

        return Response({
            'success': success
        })

    @action(detail=False, url_path='otp-status')
    def get_otp_status(self, *args, **kwargs):
        otp_good = cache.get('otp_good', False)
        return Response({
            'otp_good': otp_good
        })

    @action(detail=False, url_path='otp-required', permission_classes=[OTPRequired])
    def page_with_otp_required(self, *args, **kwargs):
        return Response({
            'success': True
        })

    @action(detail=False, url_path="edit-object", methods=['POST'], permission_classes=[OTPRequired])
    def edit_object(self, request, *args, **kwargs):
        # Логика редактирования объекта
        return Response({
            'success': True
        })
